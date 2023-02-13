import sqlite3
import docx
import matrix

def read_single_row(matrix_id):
    try:
        sqlite_connection = sqlite3.connect('db_matrix.db')
        cursor = sqlite_connection.cursor()

        sqlite_select_query = """SELECT * FROM matrix where id = ?"""
        cursor.execute(sqlite_select_query, (matrix_id, ))
        record = cursor.fetchone()
        personal_matrix = [record[1], record[2], record[3], record[4], record[5], record[6], record[7], record[8], record[9], record[10], record[11], record[12]]
        cursor.close()
        return personal_matrix
    except sqlite3.Error as error:
        print("Ошибка при работе с SQLite", error)
    finally:
        if sqlite_connection:
            sqlite_connection.close()

c_l_intro = str((read_single_row(matrix.c_l))[0])
c_l_tasks = str((read_single_row(matrix.c_l))[1])
d_l_task = str((read_single_row(matrix.d_l))[2])
d_l_minus = str((read_single_row(matrix.d_l))[3])
d_l_results = str((read_single_row(matrix.d_l))[4])
e_l = str((read_single_row(matrix.e_l))[5])
c_n_intro = str((read_single_row(matrix.c_n))[0])
c_n_tasks = str((read_single_row(matrix.c_n))[1])
d_n_task = str((read_single_row(matrix.d_n))[2])
d_n_intro = str((read_single_row(matrix.d_n))[6])
d_n_plus = str((read_single_row(matrix.d_n))[7])
d_n_conclusion = str((read_single_row(matrix.d_n))[8])
a_n = str((read_single_row(matrix.a_n))[9])
b_n = str((read_single_row(matrix.b_n))[10])
e_n = str((read_single_row(matrix.e_n))[5])

def replace_string(filename):
    document = docx.Document(filename)
    for paragraph in document.paragraphs:
        if '{ c_l_intro }' in paragraph.text:
            inline = paragraph.runs
            for i in range(len(inline)):
                if '{ c_l_intro }' in inline[i].text:
                    text = inline[i].text.replace("{ c_l_intro }", str(c_l_intro))
                    inline[i].text = text

#for paragraph in document.paragraphs:
    #paragraph.text = paragraph.text.replace('{ c_l_tasks }', str(c_l_tasks))
#for paragraph in document.paragraphs:
    #paragraph.text = paragraph.text.replace("{ d_l_task }", str(d_l_task))
#for paragraph in document.paragraphs:
    #paragraph.text = paragraph.text.replace("{ d_l_minus }", str(d_l_minus))
#for paragraph in document.paragraphs:
    #paragraph.text = paragraph.text.replace("{ d_l_results }", str(d_l_results)) 
#for paragraph in document.paragraphs:
    #paragraph.text = paragraph.text.replace("{ e_l }", str(e_l))
#for paragraph in document.paragraphs:
    #paragraph.text = paragraph.text.replace("{ c_n_intro }", str(c_n_intro))
#for paragraph in document.paragraphs:
    #paragraph.text = paragraph.text.replace("{ c_n_tasks }", str(c_n_tasks))
#for paragraph in document.paragraphs:
    #paragraph.text = paragraph.text.replace("{ d_n_task }", str(d_n_task))
#for paragraph in document.paragraphs:
    #paragraph.text = paragraph.text.replace("{ d_n_intro }", str(d_n_intro))
#for paragraph in document.paragraphs:
    #paragraph.text = paragraph.text.replace("{ d_n_plus }", str(d_n_plus))
#for paragraph in document.paragraphs:
    #paragraph.text = paragraph.text.replace("{ d_n_conclusion }", str(d_n_conclusion))
#for paragraph in document.paragraphs:
    #paragraph.text = paragraph.text.replace("{ a_n }", str(a_n))
#for paragraph in document.paragraphs:
    #paragraph.text = paragraph.text.replace("{ b_n }", str(b_n))
#for paragraph in document.paragraphs:
    #paragraph.text = paragraph.text.replace("{ e_n }", str(e_n))

    document.save("your_matrix_1.docx")
    return 1
replace_string("your_matrix_template.docx")