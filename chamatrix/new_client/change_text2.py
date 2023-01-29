import docx
import matrix
 
document = docx.Document("static/test2.docx")
 
for paragraph in document.paragraphs:
    paragraph.text = paragraph.text.replace("{{ a }}", str(matrix.a))
for paragraph in document.paragraphs:
    paragraph.text = paragraph.text.replace("{{ b }}", str(matrix.b))
 
document.save("static/test2_update.docx")